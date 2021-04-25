from __future__ import print_function
import time, json, docx, requests
import boto3
import pandas as pd
import eng_to_ipa as ipa
from docx.enum.text import WD_COLOR_INDEX

def upload_file(file_name, bucket, object_name=None):
    # If S3 object_name was not specified, use file_name
    if object_name is None:
        object_name = file_name
    s3 = boto3.client('s3', region_name = 'ap-southeast-2')
    s3.upload_file(file_name, bucket, object_name)
    print('uploaded to s3://'+bucket+'/'+object_name)
    job_uri = 's3://' + bucket + '/' + object_name
    return job_uri

def create_vocab_table(txt_list): #converts a list of vocab in a txt file (i.e., a list-like custom vocabulary) to a table with IPA etc. (https://docs.aws.amazon.com/transcribe/latest/dg/how-vocabulary.html#create-vocabulary-table)
    df = pd.read_csv(txt_list, names=['Phrase', 'IPA', 'SoundsLike', 'DisplayAs'])
    #Create a list for IPA data
    IPA = []
    for phrase in df['Phrase']:
        if ipa.isin_cmu(phrase): #if the phrase is in the CMU dictionary
            IPA.append(' '.join(ipa.convert(phrase, keep_punct=False))) #convert it to IPA with a space between each phoneme
        else:
            IPA.append('')
    df['IPA'] = IPA
    df.to_csv(txt_list[0:-4]+'_table.txt', sep='\t', index=False)

#transcripts/training_data/CX_vocab.txt
def vocab(uri, name): #this should work for a table vocab
    transcribe = boto3.client('transcribe')
    transcribe.create_vocabulary(
        LanguageCode='en-GB',
        VocabularyFileUri=uri,
        VocabularyName=name
        )
    while True:
        status = transcribe.get_vocabulary(VocabularyName=name)
        if status['VocabularyState'] in ['PENDING']:
            print("Not ready yet...")
        elif status['VocabularyState'] in ['READY']:
            print('Ready')
            break
        elif status['VocabularyState'] in ['FAILED']:
            print(status['FailureReason'])
            break
        time.sleep(20)

#vocab('s3://transcriber-audio-samples/CX_vocab', 'CX_vocab')

job_uri = 's3://transcriber-audio-samples/out_5min.wav'
def run_transcribe(name, no_of_speakers, job_uri):
    transcribe = boto3.client('transcribe')
    job_name = name
    settings = {
        'ShowSpeakerLabels':True,
        'MaxSpeakerLabels':no_of_speakers,
        #'ChannelIdentification':True #Add this feature when testing a multi-channel file
        'VocabularyName':'CX_vocab9'
    }
    transcribe.start_transcription_job(
        TranscriptionJobName=job_name,
        Media={'MediaFileUri': job_uri},
        MediaFormat='wav',
        LanguageCode='en-GB',
        Settings=settings
    )
    start = time.time()
    while True:
        status = transcribe.get_transcription_job(TranscriptionJobName=job_name)
        if status['TranscriptionJob']['TranscriptionJobStatus'] in ['COMPLETED', 'FAILED']:
            result = pd.read_json(status['TranscriptionJob']['Transcript']['TranscriptFileUri'])
            result.to_json('transcripts/json/' + job_name + '.json')
            print(status['TranscriptionJob']['Transcript']['TranscriptFileUri'])
            end = time.time()
            time_elapsed = time.strftime('%H:%M:%S', time.gmtime(end - start))
            print('this job took', time_elapsed, 'to complete')
            break
        print("Not ready yet...")
        time.sleep(20)
    return result

#json_file = 'transcripts/asrOutput.json'
def get_transcription_data(json_data):
    try:
        with open(json_data) as f:
            data = json.load(f)
    except:
        data = json_data
    transcript = {}
    para_count = 0
    speaker_data = data['results']['speaker_labels']['segments']
    word_data = data['results']['items']
    #iterate through the item number to get the start time for each change of speaker
    for speaker in speaker_data:
        para_count += 1
        speaker_label = speaker['speaker_label']
        speaker_tag = speaker_label[-1]
        speaker_start_time = float(speaker['start_time'])
        speaker_end_time = float(speaker['end_time'])
        timestamp = f"[{time.strftime('%H:%M:%S', time.gmtime(float(speaker_start_time)))}]"
        print(timestamp, 'Speaker', speaker_tag)
        #print(speaker_start_time, speaker_end_time)
        #words_per_speaker = len(speaker_data[i]['items']) #need this to set the limits for the for loop below
        words = ''
        confidence = []
        for word_info in word_data:
            if (word_info['type'] == 'punctuation'):
                words += word_info['alternatives'][0]['content']
                #confidence.append(word_info['alternatives'][0]['confidence'])
            elif (word_info['type'] == 'pronunciation') and (speaker_start_time <= float(word_info['start_time'])) and (speaker_end_time >= float(word_info['end_time'])): #start writing words at the moment the looped speaker says their first word and stopping at their last
                words += ' ' + word_info['alternatives'][0]['content']
                confidence.append(word_info['alternatives'][0]['confidence'])
            elif speaker_end_time == float(word_info['end_time']):
                words += ' ' + word_info['alternatives'][0]['content']
                confidence.append(word_info['alternatives'][0]['confidence'])
            elif speaker_end_time < float(word_info['end_time']):
                break
        print(words.lstrip('!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~')) #lstrip method is used because I can't get punctuation to be ignored in the flow of the for loop. There is no start or end time attribute, so I have nothing to match against to start the speaker dialogue. This simply strips all leading punctuation
        '''while True:
            if confidence[0] == '0.0':
                del confidence[0]
            else:
                break'''
        print(confidence) #while statement above deletes the confidence figures for the leading punctuation (all punctuation seems to have a confidence value of '0.0')
        transcript[para_count] = [f'{timestamp} Speaker {speaker_tag}', words.lstrip('!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~'), confidence]
    return transcript

def write_to_doc(transcript, doc_file=''):
    doc = docx.Document()
    for i in transcript:
        #word_count = len(transcript[i][1].split())
        #count = 0
        conf_words = '' #initialise a variable to store the words we are confident in.
        doc.add_paragraph(transcript[i][0]) #The speaker and timestamp
        #doc.add_paragraph(transcript[i][1]) #The words
        #doc.add_paragraph(transcript[i][2]) #The confidence levels for each word and punctuation
        content = doc.add_paragraph()
        for word, confidence in zip(transcript[i][1].split(), transcript[i][2]):#Iterate through the words content and confidence levels in parallel
            #count += 1
            confidence = float(confidence)
            if (confidence <= 0.85) and (confidence >=0.5) and (confidence != 0):
                #print(i, word, confidence)
                content.add_run(conf_words)
                content.add_run(' ' + word).font.highlight_color = WD_COLOR_INDEX.YELLOW
                conf_words = '' #reset the confident words list for the next time this happens.
            elif confidence < 0.5 and confidence != 0:
                content.add_run(conf_words)
                content.add_run(' ' + word).font.highlight_color = WD_COLOR_INDEX.RED
                conf_words = '' #reset the confident words list for the next time this happens.
            else:
                conf_words += ' ' + word
        content.add_run(conf_words)
    doc.save('transcripts/docs/' + doc_file)

#s3_uri = upload_file('audio_samples/out_10min.wav', 'transcriber-audio-samples', 'out_10min.wav')
s3_uri = 's3://transcriber-audio-samples/out_10min.wav'
#transcript = run_transcribe('full_test_2AU', 3, s3_uri)
#write_to_doc(get_transcription_data(transcript), 'test_AU_1.docx')